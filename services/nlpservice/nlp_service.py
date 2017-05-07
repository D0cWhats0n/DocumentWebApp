from docx import Document
from sumy.parsers.plaintext import PlaintextParser
from sumy.summarizers.sum_basic import SumBasicSummarizer as Summarizer
from sumy.nlp.stemmers import Stemmer
from sumy.nlp.tokenizers import Tokenizer

def txt_to_string(path):
    with open(path, 'r') as myfile:
            data=myfile.read().replace('\n', '')
    return "\n" + data
 
def doc_to_string(path):   
    document = Document(path)
    fullText = ""
    for para in document.paragraphs:
        fullText += "\n" + para.text
    for table in document.tables:
        for column in table.columns:
            for cell in column.cells:
                for paragraph in cell.paragraphs:
                    fullText += "\n" + paragraph.text    
    return '\n' + fullText  


def text_file_to_string(path):
    if path.endswith(".txt"):
        return txt_to_string(path)
    elif path.endswith(".docx") or path.endswith(".doc"):
        return doc_to_string(path)
    
def summarize_text(text):
    sum_text = ""
    language = "german"
    parser = PlaintextParser.from_string(text, Tokenizer(language))
    stemmer = Stemmer(language)
    summarizer = Summarizer(stemmer)
    
    for sentence in summarizer(parser.document, 7):
      for word in sentence.words:
          sum_text += word + " "
      sum_text = sum_text[:-1]    
      sum_text +=".\n"  
    return sum_text