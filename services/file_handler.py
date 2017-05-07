# -*- coding: utf-8 -*-
"""
Created on Sun Apr 30 11:14:46 2017

@author: Jonny
"""

from docx import Document

def readTableAsDataFrame(path):
    document = Document(path)
    table_df = pd.DataFrame()
    for table in document.tables:
        for index,column in enumerate(table.columns):
            column_name = "column_" + str(index)
            paragraph_np = np.empty([len(column.cells)], dtype=object)
            #print len(column.cells)
            for cell_index,cell in enumerate(column.cells):
                paragraph_string = ''
                for paragraph in cell.paragraphs:
                    paragraph_string += paragraph.text.encode('utf-8')
                paragraph_np[cell_index] = paragraph_string        
            table_df[column_name] = pd.Series(paragraph_np) 
    return table_df;